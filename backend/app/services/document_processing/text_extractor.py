"""
Text extraction service using PyMuPDF and fallback libraries.
Handles PDF, DOCX, TXT, and XLSX files with intelligent layout detection.
"""
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
import fitz  # PyMuPDF
import PyPDF2
from docx import Document as DocxDocument
import openpyxl
from loguru import logger


@dataclass
class ExtractedElement:
    """Represents a single extracted element from a document."""
    content: str
    element_type: str  # 'text', 'table', 'title', 'header', 'footer', 'list'
    page_number: int
    metadata: Dict[str, Any]
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        return {
            'content': self.content,
            'element_type': self.element_type,
            'page_number': self.page_number,
            'metadata': self.metadata
        }


@dataclass
class DocumentStructure:
    """Represents the complete structure of an extracted document."""
    elements: List[ExtractedElement]
    total_pages: int
    has_tables: bool
    has_images: bool
    metadata: Dict[str, Any]
    
    def get_elements_by_type(self, element_type: str) -> List[ExtractedElement]:
        """Filter elements by type."""
        return [e for e in self.elements if e.element_type == element_type]
    
    def get_elements_by_page(self, page_num: int) -> List[ExtractedElement]:
        """Filter elements by page number."""
        return [e for e in self.elements if e.page_number == page_num]


class TextExtractor:
    """
    Intelligent text extraction with layout awareness.
    Uses PyMuPDF and format-specific parsers.
    """
    
    def __init__(self):
        self.supported_formats = {'.pdf', '.docx', '.doc', '.txt', '.xlsx', '.xls'}
    
    def extract_document(self, file_path: Path) -> DocumentStructure:
        """
        Extract structured content from document.
        
        Args:
            file_path: Path to document file
            
        Returns:
            DocumentStructure with all extracted elements
        """
        file_ext = file_path.suffix.lower()
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        logger.info(f"Extracting document: {file_path.name}")
        
        try:
            # Use format-specific extractors
            if file_ext == '.pdf':
                return self._extract_with_pymupdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                return self._extract_docx(file_path)
            elif file_ext == '.txt':
                return self._extract_txt(file_path)
            elif file_ext in ['.xlsx', '.xls']:
                return self._extract_excel(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
        except Exception as e:
            logger.error(f"Extraction failed for {file_path.name}: {str(e)}")
            raise Exception(f"Failed to extract document: {str(e)}")
    
    def _extract_txt(self, file_path: Path) -> DocumentStructure:
        """Extract text from plain text files."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
        except Exception as e:
            logger.error(f"Error reading text file: {str(e)}")
            raise
        
        # Split into paragraphs
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        
        # If no paragraphs (single line text), split by newlines
        if not paragraphs:
            paragraphs = [line.strip() for line in content.split('\n') if line.strip()]
        
        elements = []
        for idx, para in enumerate(paragraphs, start=1):
            element = ExtractedElement(
                content=para,
                element_type="text",
                page_number=1,
                metadata={'paragraph_index': idx}
            )
            elements.append(element)
        
        logger.info(f"Extracted {len(elements)} text elements from TXT file")
        
        return DocumentStructure(
            elements=elements,
            total_pages=1,
            has_tables=False,
            has_images=False,
            metadata={'extraction_method': 'txt_parser', 'total_elements': len(elements)}
        )
    
    def _extract_docx(self, file_path: Path) -> DocumentStructure:
        """Extract text from DOCX files."""
        try:
            doc = DocxDocument(file_path)
        except Exception as e:
            logger.error(f"Error reading DOCX file: {str(e)}")
            raise
        
        elements = []
        
        # Extract paragraphs
        for idx, para in enumerate(doc.paragraphs, start=1):
            if para.text.strip():
                # Detect if it's a heading
                elem_type = "title" if para.style.name.startswith('Heading') else "text"
                
                element = ExtractedElement(
                    content=para.text.strip(),
                    element_type=elem_type,
                    page_number=1,  # DOCX doesn't have page numbers easily accessible
                    metadata={
                        'paragraph_index': idx,
                        'style': para.style.name
                    }
                )
                elements.append(element)
        
        # Extract tables
        has_tables = len(doc.tables) > 0
        for table_idx, table in enumerate(doc.tables):
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_text.append(' | '.join(row_text))
            
            element = ExtractedElement(
                content='\n'.join(table_text),
                element_type="table",
                page_number=1,
                metadata={'table_index': table_idx}
            )
            elements.append(element)
        
        logger.info(f"Extracted {len(elements)} elements from DOCX ({len(doc.tables)} tables)")
        
        return DocumentStructure(
            elements=elements,
            total_pages=1,
            has_tables=has_tables,
            has_images=False,
            metadata={
                'extraction_method': 'docx_parser',
                'total_elements': len(elements),
                'total_paragraphs': len(doc.paragraphs),
                'total_tables': len(doc.tables)
            }
        )
    
    def _extract_excel(self, file_path: Path) -> DocumentStructure:
        """Extract text from Excel files."""
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
        except Exception as e:
            logger.error(f"Error reading Excel file: {str(e)}")
            raise
        
        elements = []
        
        for sheet_idx, sheet in enumerate(workbook.worksheets):
            # Convert sheet to text representation
            rows = []
            for row in sheet.iter_rows(values_only=True):
                row_text = [str(cell) if cell is not None else '' for cell in row]
                if any(row_text):  # Skip empty rows
                    rows.append(' | '.join(row_text))
            
            if rows:
                element = ExtractedElement(
                    content='\n'.join(rows),
                    element_type="table",
                    page_number=sheet_idx + 1,
                    metadata={
                        'sheet_name': sheet.title,
                        'sheet_index': sheet_idx,
                        'rows': len(rows),
                        'columns': sheet.max_column
                    }
                )
                elements.append(element)
        
        logger.info(f"Extracted {len(elements)} sheets from Excel file")
        
        return DocumentStructure(
            elements=elements,
            total_pages=len(workbook.worksheets),
            has_tables=True,
            has_images=False,
            metadata={
                'extraction_method': 'excel_parser',
                'total_elements': len(elements),
                'total_sheets': len(workbook.worksheets)
            }
        )
    
    def _extract_with_pymupdf(self, file_path: Path) -> DocumentStructure:
        """
        Extraction using PyMuPDF.
        Effective for PDF text extraction.
        """
        try:
            doc = fitz.open(file_path)
        except Exception as e:
            logger.error(f"Error opening PDF with PyMuPDF: {str(e)}")
            raise
        
        extracted_elements = []
        has_tables = False
        
        for page_num, page in enumerate(doc, start=1):
            # Extract text blocks
            blocks = page.get_text("blocks")
            
            for block in blocks:
                # block format: (x0, y0, x1, y1, "text", block_no, block_type)
                if len(block) >= 5:
                    text = block[4].strip()
                    
                    if not text:
                        continue
                    
                    # Simple heuristic: detect tables (blocks with many numbers/pipes)
                    is_table = self._is_likely_table(text)
                    
                    if is_table:
                        has_tables = True
                    
                    extracted_elem = ExtractedElement(
                        content=text,
                        element_type="table" if is_table else "text",
                        page_number=page_num,
                        metadata={
                            'bbox': block[:4],
                            'block_type': block[6] if len(block) > 6 else None
                        }
                    )
                    
                    extracted_elements.append(extracted_elem)
        
        total_pages = len(doc)
        doc.close()
        
        logger.info(f"PyMuPDF extracted {len(extracted_elements)} elements from {total_pages} pages")
        
        return DocumentStructure(
            elements=extracted_elements,
            total_pages=total_pages,
            has_tables=has_tables,
            has_images=False,
            metadata={
                'extraction_method': 'pymupdf',
                'total_elements': len(extracted_elements)
            }
        )
    
    @staticmethod
    def _is_likely_table(text: str) -> bool:
        """
        Heuristic to detect if text block is likely a table.
        Looks for pipe characters, tabs, and numeric density.
        """
        # Count table indicators
        has_pipes = '|' in text
        has_tabs = '\t' in text
        has_multiple_lines = text.count('\n') >= 2
        
        # Calculate numeric density
        total_chars = len(text.replace(' ', '').replace('\n', ''))
        if total_chars == 0:
            return False
        
        digit_count = sum(c.isdigit() for c in text)
        numeric_density = digit_count / total_chars
        
        # Table if: has formatting AND moderate numeric density
        is_table = (has_pipes or has_tabs) and has_multiple_lines
        is_table = is_table or (numeric_density > 0.3 and has_multiple_lines)
        
        return is_table
    
    def extract_text_only(self, file_path: Path) -> str:
        """
        Extract plain text without structure (faster).
        Useful for simple text files.
        """
        file_ext = file_path.suffix.lower()
        
        if file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        
        elif file_ext == '.pdf':
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        
        elif file_ext in ['.docx', '.doc']:
            doc = DocxDocument(file_path)
            return "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        elif file_ext in ['.xlsx', '.xls']:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            text_parts = []
            for sheet in workbook.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    row_text = [str(cell) if cell is not None else '' for cell in row]
                    if any(row_text):
                        text_parts.append(' | '.join(row_text))
            return "\n".join(text_parts)
        
        else:
            raise ValueError(f"Unsupported format for text extraction: {file_ext}")


# Global instance
text_extractor = TextExtractor()

__all__ = ['TextExtractor', 'ExtractedElement', 'DocumentStructure', 'text_extractor']